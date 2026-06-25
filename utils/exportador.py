from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def separar_questoes_gabarito(texto):
    """Separa o texto das questoes do gabarito."""
    if "GABARITO:" in texto:
        partes = texto.split("GABARITO:")
        questoes = partes[0].strip()
        gabarito = "GABARITO:\n" + partes[1].strip()
    else:
        questoes = texto.strip()
        gabarito = ""
    return questoes, gabarito

def criar_documento_prova(texto_questoes):
    """Cria o documento Word com as questoes."""
    doc = Document()

    # Configurar fonte padrao
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # Adicionar questoes
    for linha in texto_questoes.split('\n'):
        linha = linha.strip()
        if not linha:
            doc.add_paragraph()
            continue
        
        paragrafo = doc.add_paragraph(linha)
        paragrafo.style = doc.styles['Normal']

    return doc

def criar_documento_gabarito(texto_gabarito):
    """Cria o documento Word com o gabarito."""
    doc = Document()

    # Configurar fonte padrao
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # Adicionar gabarito
    for linha in texto_gabarito.split('\n'):
        linha = linha.strip()
        if not linha:
            doc.add_paragraph()
            continue

        paragrafo = doc.add_paragraph(linha)
        paragrafo.style = doc.styles['Normal']

    return doc

def exportar_prova(texto_completo, caminho_prova, caminho_gabarito):
    """Exporta a prova e o gabarito para arquivos Word separados."""
    questoes, gabarito = separar_questoes_gabarito(texto_completo)

    # Criar e salvar prova
    doc_prova = criar_documento_prova(questoes)
    doc_prova.save(caminho_prova)

    # Criar e salvar gabarito (se existir)
    if gabarito:
        doc_gabarito = criar_documento_gabarito(gabarito)
        doc_gabarito.save(caminho_gabarito)
        return True
    
    return False